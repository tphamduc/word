from docx import Document
from docx.shared import Pt
from docx.text.run import Font
from docx.enum.style import WD_STYLE_TYPE
from docx.text.run import Run
from docx.text.font import ColorFormat

def ReadText(filename):
    document = Document(filename)
    text = []
    for data in document.paragraphs:
        text.append(data.text)

    for x in text:
        print(x)

# ReadText("test.docx")

def FontSize(filename):
    document = Document(filename)
    all_12pt_text = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.size == Pt(14):
                all_12pt_text.append(p.text)

    all_12pt_text_set = []
    all_12pt_text_set = set(all_12pt_text)

    for c in all_12pt_text_set:
        print(c)

def Font_Bold(Filename):
    document = Document(Filename)
    all_text_bold = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.bold:
                all_text_bold.append(p.text)

    text_bold_set = []
    text_bold_set = set(all_text_bold)
    for c in text_bold_set:
        print(c)

# Font_Bold("word.docx")

def Font_Italic(Filename):
    document = Document(Filename)
    all_text_italic = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.italic:
                all_text_italic.append(p.text)

    text_italic_set = []
    text_italic_set = set(all_text_italic)
    for c in text_italic_set:
        print(c)

# Font_Italic("word.docx")

document = Document("word.docx")
all_text_color = []
for p in document.paragraphs:
    for run in p.runs:
        print(run.font.color.rgb)



