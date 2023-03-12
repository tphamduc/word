from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.text.run import Font
from docx.enum.style import WD_STYLE_TYPE
from docx.text.run import Run
from docx.text.font import ColorFormat
from PIL import  ImageColor
from docx.section import Sections

def ReadText(filename):
    document = Document(filename)
    text = []
    for data in document.paragraphs:
        text.append(data.text)

    for x in text:
        print(x)

# ReadText("test.docx")

def FontSize(filename):
    document = Document(filename)
    all_12pt_text = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.size == Pt(14):
                all_12pt_text.append(p.text)

    all_12pt_text_set = []
    all_12pt_text_set = set(all_12pt_text)

    for c in all_12pt_text_set:
        print(c)

def Font_Bold(Filename):
    document = Document(Filename)
    all_text_bold = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.bold:
                all_text_bold.append(p.text)

    text_bold_set = []
    text_bold_set = set(all_text_bold)
    for c in text_bold_set:
        print(c)

    return c

Font_Bold("word.docx")

def Font_Italic(Filename):
    document = Document(Filename)
    all_text_italic = []
    for p in document.paragraphs:
        for run in p.runs:
            if run.font.italic:
                all_text_italic.append(p.text)

    text_italic_set = []
    text_italic_set = set(all_text_italic)
    for c in text_italic_set:
        print(c)

# Font_Italic("word.docx")


# CheckColor
def CheckColor(filename, RgbCode):
    document = Document(filename)
    all_text_color = []
    all_hex_code = []
    for p in document.paragraphs:
        for run in p.runs:
            if ImageColor.getcolor("#" + str(run.font.color.rgb), "RGB") == RgbCode:
                all_text_color.append(p.text)

    all_hex_code = set(all_text_color)
    for c in all_hex_code:
        print(c)

# CheckColor("word.docx",(213,247,223))


# CheckPaperSize
def CheckPaperSize(filename, width, height):
    document = Document(filename)
    section = document.sections[0]
    if round(section.page_height / 914400, 2) == height and round(section.page_width / 914400, 2) == width:
        print("True")
    else:
        print('False')

# CheckPaperSize("word.docx", 11.69, 8.27)

def CheckMarginPaper(filename, left, right, top, bottom):
    document = Document(filename)
    section = document.sections[0]
    if round(section.left_margin / 914400, 2) == left and round(section.right_margin / 914400, 2) == right and round(section.top_margin / 914400, 2) == top and round(section.bottom_margin / 914400, 2) == bottom:
        print("True")
    else:
        print("False")

# CheckMarginPaper("word.docx", 0.78, 0.59, 0.39, 0.39)

def CheckParaMarginCenter(filename):
    document = Document(filename)
    all_center = []
    for p in document.paragraphs:
        if p.alignment == 1:
            all_center.append(p.text)

    all_center_set = []
    all_center_set = set(all_center)
    for c in all_center_set:
        print(c)


# CheckParaMarginCenter("word.docx")

def CheckParaMarginRight(filename):
    document = Document(filename)
    all_right = []
    for p in document.paragraphs:
        if p.alignment == 2:
            all_right.append(p.text)

    all_right_set = []
    all_right_set = set(all_right)
    for c in all_right_set:
        print(c)

def CheckParaMarginJustify(filename):
    document = Document(filename)
    all_justify = []
    for p in document.paragraphs:
        if p.alignment == 3:
            all_justify.append(p.text)

    all_justify_set = []
    all_justify_set = set(all_justify)
    for c in all_justify_set:
        print(c)


def CheckLineSpacingSingle(filename):
    Rule_Single = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 0:
            Rule_Single.append(p.text)

    Rule_Single_set = []
    Rule_Single_set = set(Rule_Single)
    for c in Rule_Single_set:
        print(c)

def CheckLineOnePointFive(filename):
    Rule_One_Point_five = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 1:
            Rule_One_Point_five.append(p.text)

    Rule_One_Point_five_set = []
    Rule_One_Point_five_set = set(Rule_One_Point_five)
    for c in Rule_One_Point_five_set:
        print(c)

def CheckLineDouble(filename):
    Rule_Double = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 2:
            Rule_Double.append(p.text)

    Rule_Double_set = []
    Rule_Double_set = set(Rule_Double)
    for c in Rule_Double_set:
        print(c)

def CheckLineAtleast(filename):
    Rule_Atleast = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 3:
            Rule_Atleast.append(p.text)

    Rule_Atleast_set = []
    Rule_Atleast_set = set(Rule_Atleast)
    for c in Rule_Atleast_set:
        print(c)


def CheckLineExcatly(filename):
    Rule_Exactly = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 4:
            Rule_Exactly.append(p.text)

    Rule_Exactly_set = []
    Rule_Exactly_set = set(Rule_Exactly)
    for c in Rule_Exactly_set:
        print(c)


def CheckLineMultiple(filename):
    Rule_Multiple = []
    document = Document(filename)
    for p in document.paragraphs:
        if p.paragraph_format.line_spacing_rule == 5:
            Rule_Multiple.append(p.text)

    Rule_Multiple_set = []
    Rule_Multiple_set = set(Rule_Multiple)
    for c in Rule_Multiple_set:
        print(c)

def FirstLineIndent(filename, value):
    document = Document(filename)
    Indent = []
    for p in document.paragraphs:
        if p.paragraph_format.first_line_indent != None:
            if round(p.paragraph_format.first_line_indent / 914400, 2) == value:
                Indent.append(p.text)

    Indent_set = []
    Indent_set = set(Indent)
    for c in Indent_set:
        print(c)

# FirstLineIndent("word.docx", 0.39)


def CheckParaSpacing(filename, before, after):
    document = Document(filename)
    before = []
    after = []
    for p in document.paragraphs:
        if p.paragraph_format.space_before != None:
            if (p.paragraph_format.space_before / 12700) == before:
                before.append(p.text)
        if p.paragraph_format.space_after != None:
            if (p.paragraph_format.space_after / 12700) == after:
                after.append(p.text)

    before_set = []
    after_set = []
    before_set = set(before) 
    after_set = set(after)
    for c in before_set:
        print(c)
    for a in after_set:
        print(a)

# CheckParaSpacing("word.docx", 18.0, 10.0)
